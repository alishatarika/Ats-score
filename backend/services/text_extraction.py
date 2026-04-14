from fastapi import UploadFile
import pdfplumber
import docx
import io


def extract_text(file: UploadFile) -> str:
    """
    Extract plain text from a PDF or DOCX UploadFile.
    Returns the extracted text as a string.
    """
    filename = file.filename.lower()
    content = file.file.read()

    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.endswith(".docx"):
        return _extract_docx(content)
    else:
        return ""


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        except Exception:
            pass
    return "\n".join(text_parts)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    text_parts = []
    try:
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    except Exception:
        pass
    return "\n".join(text_parts)
